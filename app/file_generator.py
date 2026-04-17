"""
Orquesta AI — Generador de archivos premium
Genera Excel, Word y PDF de calidad profesional
"""
import io
import re
from datetime import datetime


# ─────────────────────────────────────────────────────────────────────────────
#  SYSTEM PROMPTS PARA GENERACIÓN DE ARCHIVOS
# ─────────────────────────────────────────────────────────────────────────────

FILE_SYSTEM_PROMPTS = {
    "xlsx": """Sos un experto en Excel y análisis de datos. Tu tarea es generar contenido estructurado para planillas Excel profesionales.

FORMATO DE RESPUESTA OBLIGATORIO:
Respondé ÚNICAMENTE con datos tabulares en formato CSV separado por pipes (|).
Primera línea: nombres de columnas
Siguientes líneas: datos

REGLAS:
- Generá al menos 15-20 filas de datos reales y útiles
- Usá datos concretos, no placeholders como "dato1" o "ejemplo"
- Incluí fórmulas Excel cuando sea apropiado (escribilas como texto: SUM, AVERAGE, etc.)
- Los números deben ser realistas para el contexto
- Fechas en formato DD/MM/YYYY
- Respondé SOLO con los datos, sin explicaciones ni markdown""",

    "docx": """Sos un escritor profesional experto en documentos corporativos. Tu tarea es generar contenido para documentos Word de alta calidad.

FORMATO DE RESPUESTA:
Usá estos marcadores para estructurar el documento:
# TÍTULO PRINCIPAL
## Subtítulo
### Subtítulo menor
**texto en negrita**
- item de lista
1. lista numerada
---
(separador de sección)

REGLAS:
- Generá contenido completo, profesional y detallado
- Mínimo 500 palabras de contenido real
- Usá lenguaje formal y profesional
- Incluí todos los apartados necesarios para el tipo de documento""",

    "pdf": """Sos un experto en creación de documentos PDF profesionales. Tu tarea es generar contenido estructurado y completo.

FORMATO DE RESPUESTA:
Usá estos marcadores:
# TÍTULO
## Sección
### Subsección
**negrita**
- bullets
1. numerado
---

REGLAS:
- Contenido completo y profesional
- Mínimo 400 palabras
- Estructura clara con secciones bien definidas
- Datos concretos y útiles"""
}


# ─────────────────────────────────────────────────────────────────────────────
#  GENERADOR EXCEL PREMIUM
# ─────────────────────────────────────────────────────────────────────────────

def generate_excel(ai_content: str, title: str) -> bytes:
    """Genera un archivo Excel premium con formato profesional."""
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side,
            GradientFill
        )
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import BarChart, Reference
    except ImportError:
        raise Exception("openpyxl no instalado")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:30] if title else "Datos"

    # ── Colores corporativos ──────────────────────────────────────────────────
    GREEN_DARK  = "1D9E75"
    GREEN_MID   = "5DCAA5"
    GREEN_LIGHT = "E8F8F3"
    DARK_BG     = "0D1117"
    GRAY_LIGHT  = "F5F7F5"
    GRAY_MID    = "E0E5E0"
    WHITE       = "FFFFFF"
    DARK_TEXT   = "1A1F1A"

    # ── Parsear el contenido ──────────────────────────────────────────────────
    lines = [l.strip() for l in ai_content.strip().split("\n") if l.strip()]
    
    # Detectar si es formato CSV con pipes
    data_rows = []
    headers = []
    
    for line in lines:
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                if not headers:
                    headers = cells
                else:
                    data_rows.append(cells)
        elif not headers and "," in line:
            # Intentar CSV normal
            cells = [c.strip() for c in line.split(",")]
            if not headers:
                headers = cells
            else:
                data_rows.append(cells)

    # Si no hay datos estructurados, crear estructura básica desde el texto
    if not headers:
        headers, data_rows = _parse_unstructured(ai_content, title)

    # ── Fila de título principal ───────────────────────────────────────────────
    num_cols = max(len(headers), 1)
    ws.merge_cells(f"A1:{get_column_letter(num_cols)}1")
    title_cell = ws["A1"]
    title_cell.value = f"📊 {title}"
    title_cell.font = Font(name="Calibri", size=16, bold=True, color=WHITE)
    title_cell.fill = PatternFill("solid", fgColor=GREEN_DARK)
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 35

    # ── Fila de fecha ──────────────────────────────────────────────────────────
    ws.merge_cells(f"A2:{get_column_letter(num_cols)}2")
    date_cell = ws["A2"]
    date_cell.value = f"Generado por Orquesta AI · {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    date_cell.font = Font(name="Calibri", size=9, italic=True, color="666666")
    date_cell.fill = PatternFill("solid", fgColor=GREEN_LIGHT)
    date_cell.alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 18

    # ── Fila vacía ─────────────────────────────────────────────────────────────
    ws.row_dimensions[3].height = 8

    # ── Headers ───────────────────────────────────────────────────────────────
    header_row = 4
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=header.upper())
        cell.font = Font(name="Calibri", size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=GREEN_DARK)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[header_row].height = 28

    # ── Datos ─────────────────────────────────────────────────────────────────
    for row_idx, row_data in enumerate(data_rows, header_row + 1):
        is_even = (row_idx - header_row) % 2 == 0
        row_color = GREEN_LIGHT if is_even else WHITE

        for col_idx in range(1, num_cols + 1):
            value = row_data[col_idx - 1] if col_idx <= len(row_data) else ""
            cell = ws.cell(row=row_idx, column=col_idx, value=_parse_value(value))
            cell.font = Font(name="Calibri", size=10, color=DARK_TEXT)
            cell.fill = PatternFill("solid", fgColor=row_color)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = border

            # Números y montos alineados a la derecha
            if isinstance(cell.value, (int, float)):
                cell.alignment = Alignment(horizontal="right", vertical="center")
                if col_idx > 1:
                    cell.number_format = "#,##0.00" if "." in str(value) else "#,##0"

        ws.row_dimensions[row_idx].height = 22

    # ── Totales si hay columnas numéricas ─────────────────────────────────────
    total_row = header_row + len(data_rows) + 1
    has_totals = False

    for col_idx in range(1, num_cols + 1):
        col_values = []
        for r in range(header_row + 1, total_row):
            v = ws.cell(row=r, column=col_idx).value
            if isinstance(v, (int, float)):
                col_values.append(v)

        cell = ws.cell(row=total_row, column=col_idx)
        if col_values and len(col_values) > 1:
            col_letter = get_column_letter(col_idx)
            cell.value = sum(col_values)
            cell.number_format = "#,##0.00" if any("." in str(v) for v in col_values) else "#,##0"
            has_totals = True
        elif col_idx == 1 and col_values:
            cell.value = "TOTAL"
        
        cell.font = Font(name="Calibri", size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=GREEN_MID)
        cell.alignment = Alignment(horizontal="right" if col_values else "left", vertical="center")
        cell.border = border

    if has_totals:
        ws.row_dimensions[total_row].height = 24

    # ── Ajustar anchos de columna ──────────────────────────────────────────────
    for col_idx in range(1, num_cols + 1):
        col_letter = get_column_letter(col_idx)
        max_len = len(str(headers[col_idx-1])) if col_idx <= len(headers) else 10
        for row in ws.iter_rows(min_row=header_row+1, max_row=min(header_row+len(data_rows), header_row+50)):
            cell_val = str(row[col_idx-1].value or "")
            max_len = max(max_len, min(len(cell_val), 40))
        ws.column_dimensions[col_letter].width = max(12, max_len + 4)

    # ── Freeze panes y filtros ─────────────────────────────────────────────────
    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(num_cols)}{header_row}"

    # ── Agregar hoja de resumen si hay suficientes datos ──────────────────────
    if len(data_rows) >= 5 and num_cols >= 2:
        _add_summary_sheet(wb, ws, title, headers, data_rows, header_row, 
                          GREEN_DARK, GREEN_LIGHT, WHITE, DARK_TEXT, border)

    # ── Propiedades del archivo ────────────────────────────────────────────────
    wb.properties.title = title
    wb.properties.creator = "Orquesta AI"
    wb.properties.description = f"Generado automáticamente por Orquesta AI · {datetime.now().strftime('%d/%m/%Y')}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _add_summary_sheet(wb, data_ws, title, headers, data_rows, header_row,
                       GREEN_DARK, GREEN_LIGHT, WHITE, DARK_TEXT, border):
    """Agrega una hoja de resumen/dashboard."""
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    ws2 = wb.create_sheet("📋 Resumen", 0)
    ws2.sheet_view.showGridLines = False

    # Título
    ws2.merge_cells("A1:F1")
    ws2["A1"].value = f"📊 RESUMEN — {title.upper()}"
    ws2["A1"].font = Font(name="Calibri", size=14, bold=True, color=WHITE)
    ws2["A1"].fill = PatternFill("solid", fgColor=GREEN_DARK)
    ws2["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 32

    ws2.merge_cells("A2:F2")
    ws2["A2"].value = f"Total de registros: {len(data_rows)}  ·  Columnas: {len(headers)}  ·  Generado: {datetime.now().strftime('%d/%m/%Y')}"
    ws2["A2"].font = Font(name="Calibri", size=10, italic=True, color="444444")
    ws2["A2"].fill = PatternFill("solid", fgColor=GREEN_LIGHT)
    ws2["A2"].alignment = Alignment(horizontal="center")

    # Stats de columnas numéricas
    row = 4
    ws2.cell(row=row, column=1, value="Columna").font = Font(bold=True, color=WHITE)
    ws2.cell(row=row, column=1).fill = PatternFill("solid", fgColor=GREEN_DARK)
    for col, stat in enumerate(["Mínimo", "Máximo", "Suma", "Promedio", "Registros"], 2):
        ws2.cell(row=row, column=col, value=stat).font = Font(bold=True, color=WHITE)
        ws2.cell(row=row, column=col).fill = PatternFill("solid", fgColor=GREEN_DARK)
        ws2.cell(row=row, column=col).alignment = Alignment(horizontal="center")

    row += 1
    for col_idx, header in enumerate(headers, 1):
        vals = []
        for r_data in data_rows:
            if col_idx <= len(r_data):
                v = _parse_value(r_data[col_idx - 1])
                if isinstance(v, (int, float)):
                    vals.append(v)
        if vals:
            is_even = (row % 2 == 0)
            bg = GREEN_LIGHT if is_even else WHITE
            ws2.cell(row=row, column=1, value=header).font = Font(bold=True)
            ws2.cell(row=row, column=1).fill = PatternFill("solid", fgColor=bg)
            for col, val in enumerate([min(vals), max(vals), sum(vals), sum(vals)/len(vals), len(vals)], 2):
                c = ws2.cell(row=row, column=col, value=round(val, 2))
                c.fill = PatternFill("solid", fgColor=bg)
                c.alignment = Alignment(horizontal="right")
            row += 1

    for col in range(1, 7):
        ws2.column_dimensions[get_column_letter(col)].width = 18


def _parse_value(val: str):
    """Intenta convertir un string a número o fecha."""
    if not val or val == "-":
        return val
    # Intentar número
    clean = val.replace(",", "").replace("$", "").replace("%", "").strip()
    try:
        if "." in clean:
            return float(clean)
        return int(clean)
    except:
        pass
    return val


def _parse_unstructured(content: str, title: str):
    """Crea estructura básica desde contenido no estructurado."""
    lines = [l.strip() for l in content.split("\n") if l.strip()]
    headers = ["Item", "Descripción", "Valor", "Estado"]
    data_rows = []
    for i, line in enumerate(lines[:30], 1):
        if len(line) > 5:
            data_rows.append([str(i), line[:80], "", ""])
    if not data_rows:
        data_rows = [["1", "Sin datos generados", "0", "Pendiente"]]
    return headers, data_rows


# ─────────────────────────────────────────────────────────────────────────────
#  GENERADOR WORD PREMIUM
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(ai_content: str, title: str) -> bytes:
    """Genera un archivo Word premium con formato profesional."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree
    except ImportError:
        raise Exception("python-docx no instalado")

    doc = Document()

    # ── Configurar márgenes ───────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Colores ───────────────────────────────────────────────────────────────
    GREEN = RGBColor(0x1D, 0x9E, 0x75)
    DARK  = RGBColor(0x1A, 0x1F, 0x1A)
    GRAY  = RGBColor(0x66, 0x66, 0x66)

    # ── Portada ───────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = GREEN

    doc.add_paragraph()  # espacio

    # Línea decorativa
    line_para = doc.add_paragraph("─" * 50)
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.runs[0].font.color.rgb = GREEN

    # Subtítulo con fecha
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"Generado por Orquesta AI  ·  {datetime.now().strftime('%d de %B de %Y')}")
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = GRAY

    doc.add_paragraph()  # espacio

    # ── Procesar contenido ────────────────────────────────────────────────────
    lines = ai_content.split("\n")

    for line in lines:
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # Título H1
        if line.startswith("# ") and not line.startswith("## "):
            h = doc.add_heading(line[2:], level=1)
            h.runs[0].font.color.rgb = GREEN
            h.runs[0].font.size = Pt(16)
            continue

        # Título H2
        if line.startswith("## ") and not line.startswith("### "):
            h = doc.add_heading(line[3:], level=2)
            h.runs[0].font.color.rgb = GREEN
            h.runs[0].font.size = Pt(13)
            continue

        # Título H3
        if line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            h.runs[0].font.size = Pt(11)
            continue

        # Separador
        if line.strip() == "---":
            hr = doc.add_paragraph("─" * 60)
            hr.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            hr.runs[0].font.size = Pt(8)
            continue

        # Lista con bullets
        if line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, line[2:], DARK)
            continue

        # Lista numerada
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, re.sub(r"^\d+\. ", "", line), DARK)
            continue

        # Párrafo normal con formato
        p = doc.add_paragraph()
        _add_formatted_run(p, line, DARK)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Orquesta AI  ·  La primera IA diseñada para todas las capacidades  ·  orquesta.up.railway.app")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_run(para, text: str, default_color):
    """Agrega texto con formato **negrita** al párrafo."""
    from docx.shared import Pt, RGBColor
    GREEN = RGBColor(0x1D, 0x9E, 0x75)
    
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = GREEN
        else:
            if part:
                run = para.add_run(part)
                run.font.color.rgb = default_color
                run.font.size = Pt(11)


# ─────────────────────────────────────────────────────────────────────────────
#  GENERADOR PDF PREMIUM
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(ai_content: str, title: str) -> bytes:
    """Genera un archivo PDF premium con formato profesional."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         HRFlowable, Table, TableStyle, PageBreak)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.pdfbase import pdfmetrics
    except ImportError:
        raise Exception("reportlab no instalado")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.5*cm, leftMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title=title, author="Orquesta AI"
    )

    # ── Colores ───────────────────────────────────────────────────────────────
    GREEN   = colors.HexColor("#1D9E75")
    GREEN_L = colors.HexColor("#E8F8F3")
    DARK    = colors.HexColor("#1A1F1A")
    GRAY    = colors.HexColor("#666666")
    WHITE   = colors.white

    # ── Estilos ───────────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    style_title = ParagraphStyle("OrcTitle",
        fontSize=22, fontName="Helvetica-Bold",
        textColor=GREEN, alignment=TA_CENTER, spaceAfter=6)

    style_subtitle = ParagraphStyle("OrcSub",
        fontSize=10, fontName="Helvetica-Oblique",
        textColor=GRAY, alignment=TA_CENTER, spaceAfter=20)

    style_h1 = ParagraphStyle("OrcH1",
        fontSize=15, fontName="Helvetica-Bold",
        textColor=GREEN, spaceBefore=14, spaceAfter=6,
        borderPad=4, leftIndent=0)

    style_h2 = ParagraphStyle("OrcH2",
        fontSize=12, fontName="Helvetica-Bold",
        textColor=DARK, spaceBefore=10, spaceAfter=4)

    style_h3 = ParagraphStyle("OrcH3",
        fontSize=11, fontName="Helvetica-BoldOblique",
        textColor=GRAY, spaceBefore=8, spaceAfter=3)

    style_body = ParagraphStyle("OrcBody",
        fontSize=10, fontName="Helvetica",
        textColor=DARK, leading=16, spaceAfter=6,
        alignment=TA_JUSTIFY)

    style_bullet = ParagraphStyle("OrcBullet",
        fontSize=10, fontName="Helvetica",
        textColor=DARK, leading=15, spaceAfter=3,
        leftIndent=20, bulletIndent=10)

    style_footer = ParagraphStyle("OrcFooter",
        fontSize=8, fontName="Helvetica-Oblique",
        textColor=GRAY, alignment=TA_CENTER)

    # ── Construir contenido ───────────────────────────────────────────────────
    story = []

    # Portada
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(title.upper(), style_title))
    story.append(HRFlowable(width="100%", thickness=2, color=GREEN, spaceAfter=6))
    story.append(Paragraph(
        f"Generado por <b>Orquesta AI</b>  ·  {datetime.now().strftime('%d de %B de %Y')}",
        style_subtitle))
    story.append(Spacer(1, 0.5*cm))

    # Contenido
    lines = ai_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            story.append(Spacer(1, 0.2*cm))
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            story.append(Paragraph(line[2:], style_h1))
            story.append(HRFlowable(width="40%", thickness=1, color=GREEN, spaceAfter=4))
            i += 1
            continue

        if line.startswith("## ") and not line.startswith("### "):
            story.append(Paragraph(line[3:], style_h2))
            i += 1
            continue

        if line.startswith("### "):
            story.append(Paragraph(line[4:], style_h3))
            i += 1
            continue

        if line.strip() == "---":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC"), spaceAfter=6))
            i += 1
            continue

        if line.startswith("- ") or line.startswith("• "):
            story.append(Paragraph(f"• {line[2:]}", style_bullet))
            i += 1
            continue

        if re.match(r"^\d+\. ", line):
            num = re.match(r"^(\d+)\. ", line).group(1)
            story.append(Paragraph(f"{num}. {re.sub(r'^\d+\. ', '', line)}", style_bullet))
            i += 1
            continue

        # Reemplazar **negrita** por tags HTML de ReportLab
        formatted = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", line)
        formatted = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", formatted)
        story.append(Paragraph(formatted, style_body))
        i += 1

    # Footer
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=GREEN, spaceAfter=6))
    story.append(Paragraph(
        "Orquesta AI  ·  La primera IA diseñada para todas las capacidades  ·  orquesta.up.railway.app",
        style_footer))

    # Número de página
    def add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(GRAY)
        canvas.drawRightString(A4[0] - 2.5*cm, 1.5*cm, f"Página {doc.page}")
        canvas.drawString(2.5*cm, 1.5*cm, f"Orquesta AI — {title}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return buf.getvalue()
